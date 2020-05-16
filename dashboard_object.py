from pyodbc import Error
import xlsxwriter
import excel_fotmat_modul as excel
import os
import xlrd
import pyautogui
from PIL import ImageGrab
import docx
import datetime
import time


def day_diff(cursor, table, key):
    sql_query = "SELECT DATEDIFF(DAY, (SELECT date_start FROM {0} WHERE isbn='{1}'), GETDATE()) as diff".format(
        table, key)
    try:
        cursor.execute(sql_query)
        sql_row = cursor.fetchone()
        diff = sql_row.diff
        return diff
    except Error as err:
        print(err)
        return 0


def select_price_by_key(cursor, table, key):
    sql_query = "SELECT np FROM {0} WHERE isbn='{1}'".format(table, key)
    try:
        cursor.execute(sql_query)
        sql_row = cursor.fetchone()
        price = sql_row.np
        return price
    except Error as err:
        print(err)
        return 0


class SqlCell(object):
    def __init__(self, table, key, connection, cursor):
        self.table = table
        self.key = key
        self.connection = connection
        self.cursor = cursor
        self.np = 0
        self.sp = 0
        self.ap = 0
        self.link = 0
        self.get_prices()

    def get_prices(self):
        sql_query = "SELECT np, sp, ap, link FROM {0} WHERE isbn='{1}'".format(self.table, self.key)
        try:
            self.cursor.execute(sql_query)
        except Error as err:
            print(err)
        else:
            sql_row = self.cursor.fetchone()
            if sql_row:
                self.np = sql_row.np
                self.sp = sql_row.sp
                self.ap = sql_row.ap
                self.link = sql_row.link


class PriceReport(object):
    def __init__(self, connection, cursor):
        self.connection = connection
        self.cursor = cursor
        self.rrc_block = self.get_sql_data()

    def get_sql_data(self):
        sql_query = "SELECT * FROM dbo.rrc_block"
        try:
            self.cursor.execute(sql_query)
        except Error as err:
            print(err)
            return []
        else:
            sql_data = self.cursor.fetchall()
            return sql_data


class BestReport(object):
    def __init__(self, connection, cursor):
        self.connection = connection
        self.cursor = cursor
        self.best_block = self.get_sql_data()

    def get_sql_data(self):
        sql_query = "SELECT * FROM dbo.best_block"
        try:
            self.cursor.execute(sql_query)
        except Error as err:
            print(err)
            return None
        else:
            sql_data = self.cursor.fetchall()
            return sql_data


class NewReport(object):
    def __init__(self, connection, cursor):
        self.connection = connection
        self.cursor = cursor
        self.new_block = self.get_sql_data()

    def get_sql_data(self):
        sql_query = "SELECT * FROM dbo.new_block WHERE 30>(SELECT DATEDIFF(DAY, date_start, GETDATE()))"
        try:
            self.cursor.execute(sql_query)
        except Error as err:
            print(err)
            return None
        else:
            sql_data = self.cursor.fetchall()
            return sql_data


class Report(object):
    def __init__(self, connection, cursor):
        self.connection = connection
        self.cursor = cursor
        self.RRC = PriceReport(self.connection, self.cursor)
        self.NEW = NewReport(self.connection, self.cursor)
        self.BEST = BestReport(self.connection, self.cursor)
        self.RRC_top = ['Автор', 'Наименование', 'Артикул', 'Цена']
        self.BEST_top = ['Автор', 'Наименование', 'Серия', 'Артикул', 'Цена']
        self.NEW_top = ['Автор', 'Наименование', 'Артикул', 'Серия', 'Дата', 'Дни в продаже', 'Цена']
        self.table_list = self.get_table_list()
        self.base_path = r'..\Desktop\ПАРСИНГ'

    def get_table_list(self):
        table_list = [table.table_name for table in self.cursor.tables() if 'customer' in table.table_name]
        return table_list

    def rrc_report(self, date):
        file_name = 'РРЦ {0}.xlsx'.format(date)
        path = os.path.normpath(os.path.join(self.base_path, date, file_name))
        workbook = xlsxwriter.Workbook(path)
        worksheet_np = workbook.add_worksheet('NORMAL')
        worksheet_sp = workbook.add_worksheet('SHOP')
        worksheet_ap = workbook.add_worksheet('ACTION')

        """TOP-LEFT-BLOCK"""
        tlb_col = 0
        for elem in self.RRC_top:
            worksheet_np.write(0, tlb_col, elem, excel.text_left_top_format(workbook))
            worksheet_sp.write(0, tlb_col, elem, excel.text_left_top_format(workbook))
            worksheet_ap.write(0, tlb_col, elem, excel.text_left_top_format(workbook))
            tlb_col += 1

        """TEXT-ROW-BLOCK"""
        trb_row = 1
        for elem in self.RRC.rrc_block:

            worksheet_np.write(trb_row, 0, elem.author, excel.text_row_format(workbook))
            worksheet_np.write(trb_row, 1, elem.title, excel.text_row_format(workbook))
            worksheet_np.write(trb_row, 2, elem.isbn, excel.text_row_format(workbook))
            worksheet_np.write(trb_row, 3, elem.price, excel.text_row_format(workbook))

            tb_col = 4
            for table in self.table_list:
                worksheet_np.write(0, tb_col, table, excel.text_right_top_format(workbook))
                current_cell = SqlCell(table, elem.isbn, self.connection, self.cursor)
                worksheet_np.write(trb_row, tb_col, current_cell.np,
                                   excel.price_cell_format(workbook, elem.price, current_cell.np))
                tb_col += 1

            worksheet_sp.write(trb_row, 0, elem.author, excel.text_row_format(workbook))
            worksheet_sp.write(trb_row, 1, elem.title, excel.text_row_format(workbook))
            worksheet_sp.write(trb_row, 2, elem.isbn, excel.text_row_format(workbook))
            worksheet_sp.write(trb_row, 3, elem.price, excel.text_row_format(workbook))

            tb_col = 4
            for table in self.table_list:
                worksheet_sp.write(0, tb_col, table, excel.text_right_top_format(workbook))
                current_cell = SqlCell(table, elem.isbn, self.connection, self.cursor)
                worksheet_sp.write(trb_row, tb_col, current_cell.sp,
                                   excel.price_cell_format(workbook, elem.price, current_cell.sp))
                tb_col += 1

            worksheet_ap.write(trb_row, 0, elem.author, excel.text_row_format(workbook))
            worksheet_ap.write(trb_row, 1, elem.title, excel.text_row_format(workbook))
            worksheet_ap.write(trb_row, 2, elem.isbn, excel.text_row_format(workbook))
            worksheet_ap.write(trb_row, 3, elem.price, excel.text_row_format(workbook))

            tb_col = 4
            for table in self.table_list:
                worksheet_ap.write(0, tb_col, table, excel.text_right_top_format(workbook))
                current_cell = SqlCell(table, elem.isbn, self.connection, self.cursor)
                worksheet_ap.write(trb_row, tb_col, current_cell.ap,
                                   excel.price_cell_format(workbook, elem.price, current_cell.ap))
                tb_col += 1

            trb_row += 1

        workbook.close()

    def new_report(self, date):
        file_name = 'Новинки {0}.xlsx'.format(date)
        path = os.path.normpath(os.path.join(self.base_path, date, file_name))
        workbook = xlsxwriter.Workbook(path)
        worksheet_np = workbook.add_worksheet('NORMAL')
        worksheet_sp = workbook.add_worksheet('SHOP')
        worksheet_ap = workbook.add_worksheet('ACTION')

        """TOP-LEFT-BLOCK"""
        tlb_col = 0
        for elem in self.NEW_top:
            worksheet_np.write(0, tlb_col, elem, excel.text_left_top_format(workbook))
            worksheet_sp.write(0, tlb_col, elem, excel.text_left_top_format(workbook))
            worksheet_ap.write(0, tlb_col, elem, excel.text_left_top_format(workbook))
            tlb_col += 1

        """TEXT-ROW-BLOCK"""
        trb_row = 1
        for elem in self.NEW.new_block:

            worksheet_np.write(trb_row, 0, elem.author, excel.text_row_format(workbook))
            worksheet_np.write(trb_row, 1, elem.title, excel.text_row_format(workbook))
            worksheet_np.write(trb_row, 2, elem.isbn, excel.text_row_format(workbook))
            worksheet_np.write(trb_row, 3, elem.seria, excel.text_row_format(workbook))
            worksheet_np.write(trb_row, 4, elem.date_start, excel.text_row_format(workbook))
            worksheet_np.write(trb_row, 5, day_diff(self.cursor, 'dbo.new_block', elem.isbn),
                               excel.text_row_format(workbook))
            worksheet_np.write(trb_row, 6, elem.price, excel.text_row_format(workbook))

            tb_col = 7
            for table in self.table_list:
                worksheet_np.write(0, tb_col, table, excel.text_right_top_format(workbook))
                current_cell = SqlCell(table, elem.isbn, self.connection, self.cursor)
                worksheet_np.write(trb_row, tb_col, current_cell.np,
                                   excel.price_cell_format(workbook, elem.price, current_cell.np))
                tb_col += 1

            worksheet_sp.write(trb_row, 0, elem.author, excel.text_row_format(workbook))
            worksheet_sp.write(trb_row, 1, elem.title, excel.text_row_format(workbook))
            worksheet_sp.write(trb_row, 2, elem.isbn, excel.text_row_format(workbook))
            worksheet_sp.write(trb_row, 3, elem.seria, excel.text_row_format(workbook))
            worksheet_sp.write(trb_row, 4, elem.date_start, excel.text_row_format(workbook))
            worksheet_sp.write(trb_row, 5, day_diff(self.cursor, 'dbo.new_block', elem.isbn),
                               excel.text_row_format(workbook))
            worksheet_sp.write(trb_row, 6, elem.price, excel.text_row_format(workbook))

            tb_col = 7
            for table in self.table_list:
                worksheet_sp.write(0, tb_col, table, excel.text_right_top_format(workbook))
                current_cell = SqlCell(table, elem.isbn, self.connection, self.cursor)
                worksheet_sp.write(trb_row, tb_col, current_cell.sp,
                                   excel.price_cell_format(workbook, elem.price, current_cell.sp))
                tb_col += 1

            worksheet_ap.write(trb_row, 0, elem.author, excel.text_row_format(workbook))
            worksheet_ap.write(trb_row, 1, elem.title, excel.text_row_format(workbook))
            worksheet_ap.write(trb_row, 2, elem.isbn, excel.text_row_format(workbook))
            worksheet_ap.write(trb_row, 3, elem.seria, excel.text_row_format(workbook))
            worksheet_ap.write(trb_row, 4, elem.date_start, excel.text_row_format(workbook))
            worksheet_ap.write(trb_row, 5, day_diff(self.cursor, 'dbo.new_block', elem.isbn),
                               excel.text_row_format(workbook))
            worksheet_ap.write(trb_row, 6, elem.price, excel.text_row_format(workbook))

            tb_col = 7
            for table in self.table_list:
                worksheet_ap.write(0, tb_col, table, excel.text_right_top_format(workbook))
                current_cell = SqlCell(table, elem.isbn, self.connection, self.cursor)
                worksheet_ap.write(trb_row, tb_col, current_cell.ap,
                                   excel.price_cell_format(workbook, elem.price, current_cell.ap))
                tb_col += 1

            trb_row += 1

        workbook.close()

    def best_report(self, date):
        file_name = 'Бестселлеры {0}.xlsx'.format(date)
        path = os.path.normpath(os.path.join(self.base_path, date, file_name))
        workbook = xlsxwriter.Workbook(path)
        worksheet_np = workbook.add_worksheet('NORMAL')
        worksheet_sp = workbook.add_worksheet('SHOP')
        worksheet_ap = workbook.add_worksheet('ACTION')

        """TOP-LEFT-BLOCK"""
        tlb_col = 0
        for elem in self.BEST_top:
            worksheet_np.write(0, tlb_col, elem, excel.text_left_top_format(workbook))
            worksheet_sp.write(0, tlb_col, elem, excel.text_left_top_format(workbook))
            worksheet_ap.write(0, tlb_col, elem, excel.text_left_top_format(workbook))
            tlb_col += 1

        """TEXT-ROW-BLOCK"""
        trb_row = 1
        for elem in self.BEST.best_block:

            worksheet_np.write(trb_row, 0, elem.author, excel.text_row_format(workbook))
            worksheet_np.write(trb_row, 1, elem.title, excel.text_row_format(workbook))
            worksheet_np.write(trb_row, 2, elem.seria, excel.text_row_format(workbook))
            worksheet_np.write(trb_row, 3, elem.isbn, excel.text_row_format(workbook))
            worksheet_np.write(trb_row, 4, elem.price, excel.text_row_format(workbook))

            tb_col = 5
            for table in self.table_list:
                worksheet_np.write(0, tb_col, table, excel.text_right_top_format(workbook))
                current_cell = SqlCell(table, elem.isbn, self.connection, self.cursor)
                worksheet_np.write(trb_row, tb_col, current_cell.np,
                                   excel.price_cell_format(workbook, elem.price, current_cell.np))
                tb_col += 1

            worksheet_sp.write(trb_row, 0, elem.author, excel.text_row_format(workbook))
            worksheet_sp.write(trb_row, 1, elem.title, excel.text_row_format(workbook))
            worksheet_sp.write(trb_row, 2, elem.seria, excel.text_row_format(workbook))
            worksheet_sp.write(trb_row, 3, elem.isbn, excel.text_row_format(workbook))
            worksheet_sp.write(trb_row, 4, elem.price, excel.text_row_format(workbook))

            tb_col = 5
            for table in self.table_list:
                worksheet_sp.write(0, tb_col, table, excel.text_right_top_format(workbook))
                current_cell = SqlCell(table, elem.isbn, self.connection, self.cursor)
                worksheet_sp.write(trb_row, tb_col, current_cell.sp,
                                   excel.price_cell_format(workbook, elem.price, current_cell.sp))
                tb_col += 1

            worksheet_ap.write(trb_row, 0, elem.author, excel.text_row_format(workbook))
            worksheet_ap.write(trb_row, 1, elem.title, excel.text_row_format(workbook))
            worksheet_ap.write(trb_row, 2, elem.seria, excel.text_row_format(workbook))
            worksheet_ap.write(trb_row, 3, elem.isbn, excel.text_row_format(workbook))
            worksheet_ap.write(trb_row, 4, elem.price, excel.text_row_format(workbook))

            tb_col = 5
            for table in self.table_list:
                worksheet_ap.write(0, tb_col, table, excel.text_right_top_format(workbook))
                current_cell = SqlCell(table, elem.isbn, self.connection, self.cursor)
                worksheet_ap.write(trb_row, tb_col, current_cell.ap,
                                   excel.price_cell_format(workbook, elem.price, current_cell.ap))
                tb_col += 1

            trb_row += 1

        workbook.close()

    def bad_links_file(self, date):
        file_name = 'Нарушения РРЦ {0}.xlsx'.format(date)
        for table in self.table_list:
            if table != 'customer_bookvoed' and table != 'customer_chitai':
                path = os.path.normpath(os.path.join(self.base_path, date, table, file_name))
                workbook = xlsxwriter.Workbook(path)
                worksheet = workbook.add_worksheet('Нарушения')
                worksheet.write(0, 0, 'Автор')
                worksheet.write(0, 1, 'Наименование')
                worksheet.write(0, 2, 'ISBN')
                worksheet.write(0, 3, 'РРЦ')
                worksheet.write(0, 4, 'Цена')
                worksheet.write(0, 5, 'Ссылка')
                check_count = 0
                bad_row = 1
                for elem in self.RRC.rrc_block:
                    price = SqlCell(table, elem.isbn, self.connection, self.cursor)
                    if price.np != 0 and price.np < elem.price:
                        worksheet.write(bad_row, 0, elem.author)
                        worksheet.write(bad_row, 1, elem.title)
                        worksheet.write(bad_row, 2, elem.isbn)
                        worksheet.write(bad_row, 3, elem.price)
                        worksheet.write(bad_row, 4, price.np)
                        worksheet.write(bad_row, 5, price.link)
                        worksheet.write(bad_row, 6, 'БАЗОВАЯ')
                        bad_row += 1
                        check_count += 1
                    elif price.ap != 0 and price.ap < elem.price:
                        worksheet.write(bad_row, 0, elem.author)
                        worksheet.write(bad_row, 1, elem.title)
                        worksheet.write(bad_row, 2, elem.isbn)
                        worksheet.write(bad_row, 3, elem.price)
                        worksheet.write(bad_row, 4, price.ap)
                        worksheet.write(bad_row, 5, price.link)
                        worksheet.write(bad_row, 6, 'АКЦИЯ')
                        bad_row += 1
                        check_count += 1
                if check_count > 0:
                    os.mkdir(os.path.normpath(os.path.join(self.base_path, date, table)))
                    workbook.close()

    def screenshot_modul(self):
        time.sleep(10)
        width, height = pyautogui.size()
        date = datetime.date.today().isoformat()
        global_path = os.path.normpath(os.path.join(self.base_path, date))
        list_dir = []
        for (p, d, f) in os.walk(global_path):
            if d:
                for elem in d:
                    list_dir.append(os.path.normpath(os.path.join(global_path, elem)))

        pyautogui.moveTo(width - 70, height - 25, duration=0.25)
        pyautogui.click()
        for elem in list_dir:
            str_exel = 'Нарушения РРЦ ' + date + '.xlsx'
            str_docx = 'Нарушения РРЦ ' + date + '.docx'
            excel_path = os.path.normpath(os.path.join(elem, str_exel))
            book = xlrd.open_workbook(excel_path)
            sheet = book.sheet_by_index(0)

            docx_path = os.path.normpath(os.path.join(elem, str_docx))
            document = docx.Document()
            document.save(docx_path)

            for i in range(1, sheet.nrows):
                file_name = sheet.cell(i, 2).value + '.png'
                link_name = sheet.cell(i, 5).value
                screen_path = os.path.normpath(os.path.join(elem, file_name))
                os.system('start ' + link_name)
                time.sleep(8)
                screen = ImageGrab.grab()
                screen.save(screen_path, 'PNG')
                os.system('taskkill /im chrome.exe')

                doc = docx.Document(docx_path)
                doc.add_picture(screen_path, width=docx.shared.Cm(17), height=docx.shared.Cm(7.2))
                doc.save(docx_path)
